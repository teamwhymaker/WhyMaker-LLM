from fastapi import FastAPI
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
import os
from rag import process_documents, query_rag
from dotenv import load_dotenv
import openai
from openai import OpenAI
from fastapi import File, UploadFile, Form
from typing import List
import tempfile
import json
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, UnstructuredImageLoader
from langchain.schema import HumanMessage, AIMessage, Document as TextDocument
from docx import Document as DocxDocument
from dotenv import load_dotenv
from pathlib import Path
import shutil
import tempfile as _temp
import zipfile
import subprocess
from fastapi.responses import Response, JSONResponse

load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

app = FastAPI()

ENV_FILE = os.getenv("WHYMAKER_ENV_FILE") or str(Path.home() / ".whymaker" / ".env")
ENV_DIR = os.path.dirname(ENV_FILE) or str(Path.home() / ".whymaker")

# Allow CORS for local development
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Minimal setup page to capture OpenAI key in desktop mode ---
SETUP_HTML = """
<!doctype html>
<html><head><meta charset='utf-8'><title>WhyMaker Setup</title></head>
<body style="font-family: system-ui, -apple-system, Segoe UI, Roboto, sans-serif; margin:40px;">
  <h1>WhyMaker – First‑Run Setup</h1>
  <p style="color:#666;">Version: <span id="ver">...</span></p>

  <div id="upd" style="display:none; padding:12px; border:1px solid #ddd; margin:12px 0;">
    <strong>Update available:</strong> <span id="uv"></span>
    <div style="margin-top:8px;">
      <button id="dl" style="padding:6px 10px;">Download & Install</button>
    </div>
    <div id="ustatus" style="margin-top:8px; color:#666;"></div>
  </div>

  <section style="margin-bottom:28px;">
    <h2>OpenAI</h2>
    <p>Enter your OpenAI API key:</p>
    <form id="f-openai">
      <input type="password" id="openai-key" placeholder="sk-..." style="width:360px; padding:8px;" />
      <button type="submit" style="padding:8px 12px;">Save</button>
    </form>
    <p id="m-openai" style="color:green;"></p>
  </section>

  <section>
    <h2>Google Cloud Storage (optional)</h2>
    <p>Provide your bucket name and upload a service account JSON with Storage access.</p>
    <form id="f-gcs" enctype="multipart/form-data">
      <input type="text" id="gcs-bucket" placeholder="your-bucket-name" style="width:360px; padding:8px;" />
      <input type="file" id="gcs-json" accept="application/json" style="margin-left:8px;" />
      <button type="submit" style="padding:8px 12px; margin-left:8px;">Save</button>
    </form>
    <details style="margin-top:8px;">
      <summary>Or paste JSON</summary>
      <textarea id="gcs-json-text" rows="8" cols="60" placeholder="{ \"type\": \"service_account\", ... }" style="margin-top:8px;"></textarea>
    </details>
    <p id="m-gcs" style="color:green;"></p>
  </section>

  <script>
    async function checkVersion() {
      const r = await fetch('/api/version');
      const v = await r.json();
      document.getElementById('ver').textContent = v.version;
      const c = await fetch('/api/update/check');
      const u = await c.json();
      if (u.update && u.latest) {
        document.getElementById('upd').style.display = 'block';
        document.getElementById('uv').textContent = u.latest.version;
      }
    }
    checkVersion();

    document.getElementById('f-openai').addEventListener('submit', async (e)=>{
      e.preventDefault();
      const key = (document.getElementById('openai-key').value||'').trim();
      if(!key){ alert('Please paste a key.'); return; }
      const res = await fetch('/api/settings/openai-key', { method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({ key }) });
      if(res.ok){ document.getElementById('m-openai').textContent = 'OpenAI key saved.'; }
      else{ alert('Failed to save.'); }
    });

    document.getElementById('f-gcs').addEventListener('submit', async (e)=>{
      e.preventDefault();
      const bucket = (document.getElementById('gcs-bucket').value||'').trim();
      const fileInput = document.getElementById('gcs-json');
      const text = (document.getElementById('gcs-json-text').value||'').trim();
      if(!bucket){ alert('Enter a bucket name.'); return; }
      const fd = new FormData();
      fd.append('bucket_input', bucket);
      if(fileInput.files && fileInput.files[0]){
        fd.append('file', fileInput.files[0]);
      } else if(text) {
        fd.append('creds_json', text);
      } else {
        alert('Upload or paste the service account JSON.');
        return;
      }
      const res = await fetch('/api/settings/gcs', { method:'POST', body: fd });
      if(res.ok){ document.getElementById('m-gcs').textContent = 'GCS settings saved.'; }
      else{ alert('Failed to save GCS settings.'); }
    });

    document.getElementById('dl').addEventListener('click', async ()=>{
      const st = document.getElementById('ustatus');
      st.textContent = 'Downloading update...';
      const r = await fetch('/api/update/apply', { method:'POST' });
      if (r.ok) {
        st.textContent = 'Installing and relaunching...';
      } else {
        const t = await r.text();
        st.textContent = 'Update failed: ' + t;
      }
    });
  </script>
</body></html>
"""

@app.get("/setup")
async def setup_page():
    return Response(content=SETUP_HTML, media_type="text/html")

@app.post("/api/settings/openai-key")
async def save_openai_key(payload: dict):
    key = (payload or {}).get("key", "").strip()
    if not key:
        return JSONResponse({"ok": False, "error": "missing key"}, status_code=400)
    # Persist to per-user env file
    try:
        os.makedirs(os.path.dirname(ENV_FILE), exist_ok=True)
        with open(ENV_FILE, "w", encoding="utf-8") as f:
            f.write(f"OPENAI_API_KEY={key}\n")
        # Make available in-process immediately
        os.environ["OPENAI_API_KEY"] = key
        return {"ok": True}
    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)

@app.post("/api/settings/gcs")
async def save_gcs_settings(
    bucket_input: str = Form(...),
    file: UploadFile | None = File(default=None),
    creds_json: str | None = Form(default=None),
):
    bucket_name = (bucket_input or "").strip()
    if not bucket_name:
        return JSONResponse({"ok": False, "error": "missing bucket"}, status_code=400)
    try:
        os.makedirs(ENV_DIR, exist_ok=True)
        cred_path = os.path.join(ENV_DIR, "gcp-service-account.json")
        if file and file.filename:
            content = await file.read()
            with open(cred_path, "wb") as f:
                f.write(content)
        elif creds_json:
            with open(cred_path, "w", encoding="utf-8") as f:
                f.write(creds_json)
        else:
            return JSONResponse({"ok": False, "error": "missing credentials"}, status_code=400)

        # Persist to env file
        # Merge with existing env vars if present
        existing = {}
        if os.path.exists(ENV_FILE):
            with open(ENV_FILE, "r", encoding="utf-8") as f:
                for line in f:
                    if "=" in line:
                        k, v = line.strip().split("=", 1)
                        existing[k] = v
        existing["GCS_BUCKET_NAME"] = bucket_name
        existing["GOOGLE_APPLICATION_CREDENTIALS"] = cred_path
        with open(ENV_FILE, "w", encoding="utf-8") as f:
            for k, v in existing.items():
                f.write(f"{k}={v}\n")

        # Update process env
        os.environ["GCS_BUCKET_NAME"] = bucket_name
        os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = cred_path

        # Try to initialize GCS bucket immediately
        global bucket
        try:
            from google.cloud import storage  # lazy import
            storage_client = storage.Client()
            bucket = storage_client.bucket(bucket_name)
        except Exception:
            pass

        return {"ok": True}
    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)

# --- Version & Update endpoints ---
@app.get("/api/version")
async def get_version():
    return {"version": os.getenv("WHYMAKER_VERSION", "0.0.0")}

@app.get("/api/update/check")
async def check_update():
    """Check for an update via WHYMAKER_UPDATE_URL (latest.json)."""
    import json as _json
    import urllib.request
    latest_url = os.getenv("WHYMAKER_UPDATE_URL", "").strip()
    current = os.getenv("WHYMAKER_VERSION", "0.0.0")
    if not latest_url:
        return {"update": False, "reason": "no url", "current": current}
    try:
        with urllib.request.urlopen(latest_url, timeout=8) as resp:
            data = _json.loads(resp.read().decode("utf-8"))
        latest_ver = data.get("version", "0.0.0")
        return {"update": latest_ver != current, "current": current, "latest": data}
    except Exception as e:
        return {"update": False, "error": str(e), "current": current}

@app.post("/api/update/apply")
async def apply_update():
    """Download zip, extract, replace app bundle, relaunch."""
    import json as _json
    import urllib.request
    latest_url = os.getenv("WHYMAKER_UPDATE_URL", "").strip()
    if not latest_url:
        return JSONResponse({"ok": False, "error": "Missing WHYMAKER_UPDATE_URL"}, status_code=400)
    try:
        with urllib.request.urlopen(latest_url, timeout=10) as resp:
            cfg = _json.loads(resp.read().decode("utf-8"))
        zip_url = cfg.get("url")
        if not zip_url:
            return JSONResponse({"ok": False, "error": "latest.json missing url"}, status_code=400)

        # Download zip to temp
        tmpdir = _temp.mkdtemp(prefix="whymaker_upd_")
        zip_path = os.path.join(tmpdir, "update.zip")
        with urllib.request.urlopen(zip_url, timeout=60) as resp, open(zip_path, "wb") as out:
            shutil.copyfileobj(resp, out)

        # Extract
        extract_dir = os.path.join(tmpdir, "unzipped")
        os.makedirs(extract_dir, exist_ok=True)
        with zipfile.ZipFile(zip_path, 'r') as zf:
            zf.extractall(extract_dir)

        # Find .app in extracted content
        new_app = None
        for root, dirs, files in os.walk(extract_dir):
            for d in dirs:
                if d.endswith('.app'):
                    new_app = os.path.join(root, d)
                    break
            if new_app:
                break
        if not new_app:
            return JSONResponse({"ok": False, "error": "No .app found in update"}, status_code=400)

        # Determine current app path
        # When bundled, sys.argv[0] is the full path to the Mach-O; app bundle is two dirs up
        import sys
        exe_path = sys.argv[0]
        bundle_path = Path(exe_path).resolve().parent.parent
        if not str(bundle_path).endswith('.app'):
            return JSONResponse({"ok": False, "error": "Not running from .app bundle"}, status_code=400)

        # Write an installer script to replace and relaunch (avoids locking files)
        installer = os.path.join(tmpdir, "install.sh")
        with open(installer, "w", encoding="utf-8") as f:
            f.write(f"""#!/bin/bash
set -e
TARGET=\"{bundle_path}\"
SRC=\"{new_app}\"
sleep 1
osascript -e 'tell application "System Events" to tell process "WhyMaker" to quit' || true
sleep 1
rm -rf \"$TARGET\"
cp -R \"$SRC\" \"$TARGET\"
xattr -dr com.apple.quarantine \"$TARGET\" || true
open \"$TARGET\"
""")
        os.chmod(installer, 0o755)
        subprocess.Popen([installer])
        return {"ok": True}
    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)

# --- Google Cloud Storage setup (optional) ---
DESKTOP_MODE = os.getenv("WHYMAKER_DESKTOP", "0") == "1"
GCS_BUCKET_NAME = os.getenv("GCS_BUCKET_NAME")
GOOGLE_APPLICATION_CREDENTIALS = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
bucket = None

UPLOAD_DIR = os.getenv("WHYMAKER_UPLOAD_DIR", "uploads")
# Ensure local folder exists
os.makedirs(UPLOAD_DIR, exist_ok=True)

if GCS_BUCKET_NAME and GOOGLE_APPLICATION_CREDENTIALS:
    try:
        from google.cloud import storage  # lazy import only when needed
        storage_client = storage.Client()
        bucket = storage_client.bucket(GCS_BUCKET_NAME)
        # Download any existing files from GCS → local UPLOAD_DIR
        for blob in bucket.list_blobs():
            local_path = os.path.join(UPLOAD_DIR, blob.name)
            os.makedirs(os.path.dirname(local_path), exist_ok=True)
            blob.download_to_filename(local_path)
    except Exception as e:
        print(f"GCS disabled (init failed): {e}")
        bucket = None
else:
    print("GCS disabled (missing credentials); using local uploads only.")

# --- Ingest all existing docs on startup ---
@app.on_event("startup")
async def on_startup():
    process_documents(folder_path=UPLOAD_DIR)

class Query(BaseModel):
    question: str
    chat_history: list = []
    model: str = "gpt-3.5-turbo"

@app.post("/api/chat")
async def chat(
    question: str = Form(...),
    chat_history: str = Form("[]"),
    model: str = Form("gpt-3.5-turbo"),
    files: List[UploadFile] = File([]),
):
    # 1) Parse and convert history
    history_list = json.loads(chat_history)
    converted_history = []
    for msg in history_list:
        if msg["role"] == "user":
            converted_history.append(HumanMessage(content=msg["content"]))
        else:
            converted_history.append(AIMessage(content=msg["content"]))

    # 2) Load & split all uploaded files in-memory
    ephemeral_chunks = []
    if files:
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        for upload in files:
            # write to temp file
            suffix = os.path.splitext(upload.filename)[1]
            tf = tempfile.NamedTemporaryFile(suffix=suffix, delete=False)
            tf.write(await upload.read())
            tf.flush()
            tf.close()

            # load docs just like in process_documents
            if suffix.lower() == ".pdf":
                docs = PyPDFLoader(tf.name).load()
            elif suffix.lower() in [".jpg", ".jpeg", ".png", ".svg"]:
                docs = UnstructuredImageLoader(tf.name).load()
            elif suffix.lower() == ".docx":
                # simplified in‐memory DOCX loader
                temp_doc = DocxDocument(tf.name)
                docs = [TextDocument(page_content=p.text, metadata={"source": upload.filename})
                        for p in temp_doc.paragraphs if p.text.strip()]
            else:
                # fallback: treat as plain text
                text = tf.read().decode("utf-8", errors="ignore")
                docs = [TextDocument(page_content=text, metadata={"source": upload.filename})]

            os.unlink(tf.name)
            ephemeral_chunks.extend(splitter.split_documents(docs))

    # 3) Run the RAG query, passing these ephemeral chunks
    answer, _ = query_rag(
        question,
        converted_history,
        model_name=model,
        extra_docs=ephemeral_chunks,
    )

    # 4) Generate a chat title on the first user message
    if len(history_list) == 0:
        try:
            title_prompt = (
                "Summarize the following user question into a 3-5 word title. "
                f"Be concise and representative of the main topic.\n\n"
                f"Question: '{question}'"
            )
            response = client.chat.completions.create(
                model="gpt-4.1-nano",
                messages=[
                    {"role": "system", "content": "You are a title generator for WhyMaker."},
                    {"role": "user", "content": title_prompt},
                ],
                max_tokens=15,
                temperature=0.2,
            )
            if response and response.choices:
                title = response.choices[0].message.content.strip().replace('"', "")
            else:
                title = question[:30] + "..."
        except Exception as e:
            print(f"Title generation failed: {e}")
            title = question[:30] + "..."
    else:
        title = None

    return {"answer": answer, "title": title}

@app.post("/api/upload")
async def upload(files: List[UploadFile] = File(...)):
    """
    Save each uploaded file to both GCS and local disk,
    then re-run process_documents() to ingest them.
    """
    saved = []
    for upload in files:
        filename = upload.filename
        # If GCS configured, upload there
        if bucket is not None:
            blob = bucket.blob(filename)
            blob.upload_from_file(await upload.read(), content_type=upload.content_type)
            # Also mirror locally
            local_path = os.path.join(UPLOAD_DIR, filename)
            os.makedirs(os.path.dirname(local_path), exist_ok=True)
            blob.download_to_filename(local_path)
        else:
            # Local-only save
            local_path = os.path.join(UPLOAD_DIR, filename)
            os.makedirs(os.path.dirname(local_path), exist_ok=True)
            with open(local_path, "wb") as f:
                f.write(await upload.read())
        saved.append(filename)

    # Re-process only the uploads folder so these docs become queryable
    process_documents(folder_path=UPLOAD_DIR)
    return {"status": "success", "uploaded": saved} 